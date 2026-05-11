"""
Document endpoints with integrated ML models (BERT + DistilBART)
"""

import os
import re
import io
import uuid
import time
from datetime import datetime, timedelta
from typing import Optional

import numpy as np
import torch
import docx
import PyPDF2
from fastapi import APIRouter, Depends, HTTPException, status, File, UploadFile, Form
from fastapi.responses import JSONResponse
from pydantic import BaseModel
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.future import select
from sqlalchemy import desc

from transformers import (
    BertTokenizer,
    BertForSequenceClassification,
    BartTokenizer,
    BartForConditionalGeneration,
)

try:
    from peft import PeftModel

    PEFT_AVAILABLE = True
except ImportError:
    PEFT_AVAILABLE = False

from app.database import get_db
from app.models import User, Document
from app.schemas import (
    DocumentResponse,
    DocumentListResponse,
    SummaryResponse,
    DeleteResponse,
)
from app.auth import get_current_user

router = APIRouter(prefix="/documents", tags=["documents"])

# ===== CONFIG =====

_HERE = os.path.dirname(os.path.abspath(__file__))  # .../backend/app/routers/
MODEL_DIR = os.path.normpath(os.path.join(_HERE, "..", "..", "..", "model"))
CLASSIFIER_DIR = os.path.join(MODEL_DIR, "classifier_model")
DISTILBART_DIR = os.path.join(MODEL_DIR, "distilbart-summarizer-finetuned")
DISTILBART_BASE = "sshleifer/distilbart-cnn-12-6"

ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt"}
ALLOWED_MIME_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "text/plain",
}
DEVICE = "cuda" if torch.cuda.is_available() else "cpu"
BERT_MAX_LEN = 128
MAX_INPUT_LEN = 512
MAX_TARGET_LEN = 128
LABEL_NAMES = ["notice", "feedback", "complaint"]

# ===== FILE VALIDATION =====


def validate_file(file: UploadFile) -> None:
    # Check file extension
    filename = file.filename or ""
    if not any(filename.lower().endswith(ext) for ext in ALLOWED_EXTENSIONS):
        raise HTTPException(
            status_code=400, detail="Unsupported file type. Allowed: PDF, DOCX, TXT"
        )

    # Check MIME type
    if file.content_type not in ALLOWED_MIME_TYPES:
        raise HTTPException(status_code=400, detail="Invalid file content type")

    # Basic content validation (read first few bytes)
    content = file.file.read(10)
    file.file.seek(0)  # Reset file pointer
    if filename.lower().endswith(".pdf") and not content.startswith(b"%PDF"):
        raise HTTPException(status_code=400, detail="Invalid PDF file")
    elif filename.lower().endswith(".docx") and not content.startswith(b"PK"):
        raise HTTPException(status_code=400, detail="Invalid DOCX file")
    # For TXT, no specific header


# ===== MODEL LOADING (runs once at import time) =====

DISTILBART_GEN = {
    "notice": {
        "max_new_tokens": MAX_TARGET_LEN,
        "num_beams": 4,
        "length_penalty": 1.2,
        "no_repeat_ngram_size": 4,
        "repetition_penalty": 1.3,
    },
    "feedback": {
        "max_new_tokens": MAX_TARGET_LEN,
        "num_beams": 4,
        "length_penalty": 1.0,
        "no_repeat_ngram_size": 4,
        "repetition_penalty": 1.3,
    },
    "complaint": {
        "max_new_tokens": MAX_TARGET_LEN,
        "num_beams": 4,
        "length_penalty": 1.0,
        "no_repeat_ngram_size": 4,
        "repetition_penalty": 1.2,
    },
}

# ===== MODEL LOADING (runs once at import time) =====

print("🔧 Loading BERT classifier...")
c_tok = BertTokenizer.from_pretrained(CLASSIFIER_DIR, local_files_only=True)
c_model = BertForSequenceClassification.from_pretrained(
    CLASSIFIER_DIR, local_files_only=True
).to(DEVICE)
c_model.eval()

print("🔧 Loading DistilBART summarizer...")
if not PEFT_AVAILABLE:
    raise RuntimeError("peft package is required. Run: pip install peft")

s_tok = BartTokenizer.from_pretrained(DISTILBART_DIR, local_files_only=True)
_base = BartForConditionalGeneration.from_pretrained(DISTILBART_BASE)
s_model = PeftModel.from_pretrained(_base, DISTILBART_DIR)
s_model = s_model.to(DEVICE)
s_model.eval()

print(f"✅ Models ready! Running on {DEVICE.upper()}")

# ===== TEXT HELPERS (from your upload.py) =====


def extract_text(file_bytes: bytes, extension: str) -> str:
    """Extract plain text from TXT, DOCX, or PDF bytes."""
    if extension == ".txt":
        return file_bytes.decode("utf-8")
    elif extension == ".docx":
        doc = docx.Document(io.BytesIO(file_bytes))
        return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    elif extension == ".pdf":
        reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        return "".join([page.extract_text() or "" for page in reader.pages])
    raise ValueError(f"Unsupported extension: {extension}")


def normalize_input(text: str) -> str:
    """Fix common input spacing issues."""
    titles = ["Dr", "Mr", "Mrs", "Ms", "Prof", "St", "Jr", "Sr", "Rev", "Lt", "Col"]
    for t in titles:
        text = re.sub(rf"\b{t}\.(?=[A-Z])", f"{t}. ", text)
    text = re.sub(r"([a-z])\.([A-Z])", r"\1. \2", text)
    text = re.sub(r",(?=[^\s])", ", ", text)
    text = re.sub(r" {2,}", " ", text)
    return text.strip()


def format_distilbart_input(label: str, text: str) -> str:
    """Prefix text with label token for fine-tuned DistilBART."""
    return f"summarize {label.upper()}: {normalize_input(text)}"


def clean_filler(text: str) -> str:
    """Strip boilerplate openers."""
    patterns = [
        r"^(?:dear\s+\w+[\s,]+)+",
        r"^greetings from [^.!]+[.!]\s*",
        r"^(?:we\s+(?:are\s+)?(?:pleased|excited|happy|thrilled|delighted)\s+to\s+(?:inform|announce|share|invite)[^,]*,?\s*)",
        r"^(?:this\s+is\s+to\s+(?:inform|notify)\s+(?:you\s+)?(?:all\s+)?(?:that\s+)?)",
        r"^(?:please\s+note\s+that\s+)",
        r"^(?:i\s+hope\s+this\s+(?:email|message)\s+finds\s+you[^.]+\.\s*)",
    ]
    for p in patterns:
        text = re.sub(p, "", text, flags=re.IGNORECASE).strip()
    return text


def keep_first_sentence(text: str) -> str:
    """Truncate to first complete sentence."""
    abbrevs = [
        "Dr",
        "Mr",
        "Mrs",
        "Ms",
        "Prof",
        "St",
        "Jr",
        "Sr",
        "vs",
        "etc",
        "approx",
        "Dept",
        "Govt",
        "No",
        "Vol",
        "Jan",
        "Feb",
        "Mar",
        "Apr",
        "Jun",
        "Jul",
        "Aug",
        "Sep",
        "Oct",
        "Nov",
        "Dec",
        "Ph",
        "Rev",
        "Lt",
        "Col",
        "Sgt",
        "Pvt",
        "Capt",
        "Gen",
        "Eng",
        "Asst",
        "Assoc",
    ]
    protected = text
    for ab in abbrevs:
        protected = re.sub(rf"\b{ab}\.", f"{ab}<DOT>", protected)

    m = re.search(r"(.+?[.!?])\s+[A-Z]", protected)
    if m:
        return m.group(1).replace("<DOT>", ".").strip()

    restored = protected.replace("<DOT>", ".")
    if re.search(r"[.!?]$", restored.strip()):
        return restored.strip()
    return restored.strip() + "."


def fix_punctuation(text: str) -> str:
    """Normalise spacing around punctuation."""
    text = re.sub(r"(\w)\s+[–\-]\s+(\w)", r"\1-\2", text)
    text = re.sub(r"\s+([.,!?;:])", r"\1", text)
    text = re.sub(r"([.,!?;:])([A-Za-z])", r"\1 \2", text)
    text = re.sub(r"\.\.+", ".", text)
    sentences = re.split(r"(?<=[.!?])\s+", text.strip())
    text = " ".join(s[0].upper() + s[1:] if s else s for s in sentences)
    if text and text[-1] not in ".!?":
        text += "."
    return text.strip()


def enforce_label_format(text: str, label: str) -> str:
    """Remove model output artifacts."""
    text = re.sub(r"^summarize\s+\w+\s*:\s*", "", text, flags=re.IGNORECASE).strip()
    text = re.sub(r"\bsummarize\b", "", text, flags=re.IGNORECASE).strip()

    if label in ("feedback", "complaint"):
        text = re.sub(
            r"^(?:Student\s+(?:complains|suggests|requests|appreciates|recommends|mentions|notes|feels|reports)"
            r"|Writer\s+(?:appreciates|suggests|complains|mentions))\s+",
            "",
            text,
            flags=re.IGNORECASE,
        ).strip()
        if text:
            text = text[0].upper() + text[1:]

    return text.strip()


def extract_notice_details(text: str) -> tuple[str, str, str]:
    """Extract date, time, and venue from notice text."""
    date_str = time_str = venue_str = ""
    for line in text.splitlines():
        s, lo = line.strip(), line.strip().lower()
        if lo.startswith("date") and not date_str:
            date_str = re.sub(r"^date\s*[:\-]?\s*", "", s, flags=re.IGNORECASE).strip()
        if lo.startswith("time") and not time_str:
            time_str = re.sub(r"^time\s*[:\-]?\s*", "", s, flags=re.IGNORECASE).strip()
        if (lo.startswith("venue") or lo.startswith("location")) and not venue_str:
            venue_str = re.sub(
                r"^(?:venue|location)\s*[:\-]?\s*", "", s, flags=re.IGNORECASE
            ).strip()

    if not time_str:
        m = re.search(
            r"\b\d{1,2}[:.]\d{2}\s*(?:am|pm)?\s*(?:–|-|to)\s*\d{1,2}[:.]\d{2}\s*(?:am|pm)?\b"
            r"|\b\d{1,2}[:.]\d{2}\s*(?:am|pm)?\b|\b\d{1,2}\s*(?:am|pm)\b",
            text,
            re.IGNORECASE,
        )
        if m:
            time_str = m.group()

    if not date_str:
        m = re.search(
            r"\b(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\s+\d{1,2}(?:\s*[–\-]\s*\d{1,2})?,?\s*\d{2,4}\b"
            r"|\b\d{1,2}(?:st|nd|rd|th)?\s+(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\b"
            r"|\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b",
            text,
            re.IGNORECASE,
        )
        if m:
            date_str = m.group()

    return date_str.strip(), time_str.strip(), venue_str.strip()


def extract_online_mode(text: str) -> bool:
    return any(
        re.search(p, text, re.IGNORECASE)
        for p in [
            r"\bonline\b",
            r"\bvirtual\b",
            r"\bzoom\b",
            r"\bteams\b",
            r"\bgoogle meet\b",
            r"\bwebinar\b",
            r"\blivestream\b",
            r"\bremote\b",
        ]
    )


def extract_conductor(text: str) -> str:
    m = re.search(
        r"(?:conducted|led|facilitated|organized|presented)\s+by\s+"
        r"((?:Dr|Mr|Mrs|Ms|Prof)\.?\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)",
        text,
        re.IGNORECASE,
    )
    return m.group(1).strip() if m else ""


def append_notice_fields(
    summary: str, text: str, date_str: str, time_str: str, venue_str: str
) -> str:
    """Re-attach date/time/venue/conductor details."""
    is_online = extract_online_mode(text)
    conductor = extract_conductor(text)
    parts = []
    if is_online:
        if "online" not in summary.lower():
            parts.append("conducted online")
    elif venue_str and venue_str.lower() not in summary.lower():
        parts.append(f"at {venue_str}")
    if date_str and date_str.lower() not in summary.lower():
        parts.append(f"on {date_str}")
    if time_str and time_str.lower() not in summary.lower():
        parts.append(f"at {time_str}")
    if conductor and conductor.lower() not in summary.lower():
        parts.append(f"conducted by {conductor}")
    if parts:
        summary = summary.rstrip(".") + ", " + ", ".join(parts) + "."
    return summary


# ===== INFERENCE =====


def classify(text: str) -> tuple[str, float, dict]:
    """Run BERT classifier."""
    inputs = c_tok(
        text,
        return_tensors="pt",
        truncation=True,
        padding=True,
        max_length=BERT_MAX_LEN,
    ).to(DEVICE)
    with torch.no_grad():
        probs = torch.softmax(c_model(**inputs).logits, dim=-1).squeeze().cpu().numpy()

    pred_idx = int(np.argmax(probs))
    label = LABEL_NAMES[pred_idx]
    confidence = float(probs[pred_idx])
    all_scores = {LABEL_NAMES[i]: float(probs[i]) for i in range(len(LABEL_NAMES))}
    return label, confidence, all_scores


def summarize(text: str, label: str) -> tuple[str, float]:
    """Run DistilBART with label-specific config."""
    inp_text = format_distilbart_input(label, text)
    inputs = s_tok(
        inp_text,
        return_tensors="pt",
        max_length=MAX_INPUT_LEN,
        truncation=True,
    ).to(DEVICE)

    cfg = DISTILBART_GEN[label]
    t0 = time.time()
    with torch.no_grad():
        ids = s_model.generate(
            **inputs,
            max_new_tokens=cfg["max_new_tokens"],
            num_beams=cfg["num_beams"],
            length_penalty=cfg["length_penalty"],
            no_repeat_ngram_size=cfg["no_repeat_ngram_size"],
            repetition_penalty=cfg["repetition_penalty"],
            early_stopping=True,
        )
    elapsed = time.time() - t0

    summary = s_tok.decode(ids[0], skip_special_tokens=True)

    # Post-processing
    summary = clean_filler(summary)
    summary = keep_first_sentence(summary)
    summary = enforce_label_format(summary, label)

    if label in ("complaint", "feedback"):
        words = summary.split()
        if len(words) > 30:
            summary = " ".join(words[:30]).rstrip(".,;:") + "."

    summary = fix_punctuation(summary)

    # Fallback if too short
    if len(summary.split()) <= 3:
        fallback = keep_first_sentence(normalize_input(text))
        if label in ("complaint", "feedback"):
            words = fallback.split()
            if len(words) > 30:
                fallback = " ".join(words[:30]).rstrip(".,;:") + "."
        summary = fix_punctuation(fallback)

    # For notices: re-attach details
    if label == "notice":
        date_str, time_str, venue_str = extract_notice_details(text)
        summary = append_notice_fields(summary, text, date_str, time_str, venue_str)

    return summary, elapsed


# ===== REQUEST SCHEMA =====


class TextPayload(BaseModel):
    text: str


# ===== ENDPOINTS =====


@router.post("/upload", response_model=DocumentResponse)
async def upload_document(
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """
    Upload a file (PDF, DOCX, TXT) for categorization and summarization.

    Protected endpoint - requires authentication token.
    """
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file provided")

    # Validate file type and content
    validate_file(file)

    file_bytes = await file.read()

    # Extract text
    try:
        text = extract_text(file_bytes, extension)
    except Exception as e:
        raise HTTPException(status_code=422, detail=f"Could not extract text: {str(e)}")

    if not text.strip():
        raise HTTPException(
            status_code=422, detail="Document appears empty or unreadable"
        )

    if len(text.split()) < 10:
        raise HTTPException(
            status_code=422, detail="Document too short (need at least 10 words)"
        )

    # Run models
    try:
        label, confidence, all_scores = classify(text)
        summary, inference_time = summarize(text, label)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Model inference failed: {str(e)}")

    # Save to database
    unique_name = f"{uuid.uuid4()}{extension}"
    file_path = f"uploads/{current_user.id}/{unique_name}"
    os.makedirs(os.path.dirname(file_path), exist_ok=True)

    try:
        with open(file_path, "wb") as buffer:
            buffer.write(file_bytes)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to save file: {str(e)}")

    new_doc = Document(
        user_id=current_user.id,
        original_content=text,
        file_path=file_path,
        category=label,
        confidence_score=confidence,
        summary=summary,
        expires_at=datetime.utcnow() + timedelta(days=30),
    )

    db.add(new_doc)
    await db.commit()
    await db.refresh(new_doc)

    return DocumentResponse(
        id=new_doc.id,
        original_content=new_doc.original_content,
        category=new_doc.category,
        confidence_score=new_doc.confidence_score,
        summary=new_doc.summary,
        created_at=new_doc.created_at,
        all_scores=all_scores,
    )


@router.post("/text")
async def submit_text(
    payload: TextPayload,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """
    Process raw text input through classification and summarization.

    Protected endpoint - requires authentication token.
    """
    text = payload.text.strip()

    if not text:
        raise HTTPException(status_code=422, detail="Text cannot be empty")

    if len(text.split()) < 10:
        raise HTTPException(
            status_code=422, detail="Text too short (need at least 10 words)"
        )

    # Run models
    try:
        label, confidence, all_scores = classify(text)
        summary, inference_time = summarize(text, label)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Model inference failed: {str(e)}")

    # Save to database (no file)
    new_doc = Document(
        user_id=current_user.id,
        original_content=text,
        file_path=None,  # No file uploaded
        category=label,
        confidence_score=confidence,
        summary=summary,
        expires_at=datetime.utcnow() + timedelta(days=30),
    )

    db.add(new_doc)
    await db.commit()
    await db.refresh(new_doc)

    return DocumentResponse(
        id=new_doc.id,
        original_content=new_doc.original_content,
        category=new_doc.category,
        confidence_score=new_doc.confidence_score,
        summary=new_doc.summary,
        created_at=new_doc.created_at,
        all_scores=all_scores,
    )


@router.get("/", response_model=list[DocumentListResponse])
async def list_documents(
    category: str = None,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """
    Get user's documents with optional category filter.

    Protected endpoint - requires authentication token.
    """
    stmt = select(Document).where(Document.user_id == current_user.id)

    if category:
        valid_categories = ["notice", "feedback", "complaint"]
        if category not in valid_categories:
            raise HTTPException(
                status_code=400,
                detail=f"Invalid category. Must be one of: {', '.join(valid_categories)}",
            )
        stmt = stmt.where(Document.category == category)

    stmt = stmt.order_by(desc(Document.created_at))

    result = await db.execute(stmt)
    documents = result.scalars().all()

    return [DocumentListResponse.model_validate(doc) for doc in documents]


@router.get("/{doc_id}", response_model=DocumentResponse)
async def get_document(
    doc_id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Get a specific document by ID."""
    stmt = select(Document).where(
        (Document.id == doc_id) & (Document.user_id == current_user.id)
    )
    result = await db.execute(stmt)
    doc = result.scalars().first()

    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    return DocumentResponse.model_validate(doc)


@router.delete("/{doc_id}", response_model=DeleteResponse)
async def delete_document(
    doc_id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Delete a document."""
    stmt = select(Document).where(
        (Document.id == doc_id) & (Document.user_id == current_user.id)
    )
    result = await db.execute(stmt)
    doc = result.scalars().first()

    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    # Delete file if exists
    if doc.file_path and os.path.exists(doc.file_path):
        os.remove(doc.file_path)

    # Delete from database
    await db.delete(doc)
    await db.commit()

    return DeleteResponse(message="Document deleted successfully")
